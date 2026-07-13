"""Extract plain text from downloaded / unzipped files and catalogue it.

Reads each file recorded in the ``files`` table that is present on disk, extracts
its text (PDF, DOCX, XLSX, and any plain-text / markup format), writes the text to
a ``.txt`` file under a SEPARATE output root, and records the location in the
``extracted_texts`` table. Binary, media and unsupported formats are skipped.

Only libraries already present in the project venv are used: PyMuPDF (fitz),
python-docx, openpyxl, lxml.
"""
import logging
from pathlib import Path
from typing import Optional

from pipeline.database import QDArchDatabase
from pipeline.downloader import _sanitize_path_part

log = logging.getLogger("pipeline.text_extractor")

# Rich formats with dedicated parsers.
PDF_EXTS = {"pdf"}
DOCX_EXTS = {"docx"}
XLSX_EXTS = {"xlsx", "xlsm"}
# Markup we strip down to visible text.
MARKUP_EXTS = {"html", "htm", "xml"}
# Formats that are already plain text — decoded as-is.
PLAIN_EXTS = {
    "txt", "md", "markdown", "rst", "csv", "tsv", "tab", "json", "jsonl",
    "ndjson", "geojson", "yaml", "yml", "log", "tex", "bib", "ris", "srt",
    "vtt", "sql", "r", "rmd", "qmd", "py", "ipynb", "do", "sps", "sas",
    "js", "css", "do", "m", "csl", "cff", "toml", "ini", "cfg", "dat",
}
# Known-binary / media / project formats we never try to read as text.
SKIP_EXTS = {
    # archives
    "zip", "rar", "7z", "gz", "tar", "tgz", "bz2", "xz",
    # media
    "mp4", "m4v", "mov", "avi", "mkv", "wmv", "flv", "webm", "mpg", "mpeg",
    "jpg", "jpeg", "png", "gif", "bmp", "tif", "tiff", "svg", "webp", "heic", "ico",
    "mp3", "wav", "flac", "aac", "ogg", "oga", "m4a", "wma", "aiff",
    # binary data / stats / qda / office-binary we can't parse with installed libs
    "mat", "sav", "dta", "pkl", "parquet", "feather", "h5", "hdf5", "npy", "npz",
    "gpkg", "shp", "grd", "gri", "nc", "rdata", "rds",
    "doc", "ppt", "pptx", "odt", "ods", "odp", "rtf", "xls",
    "qdpx", "nvpx", "nvp", "atlproj", "atlproj9", "atlproj23", "mx", "mx20",
    "mx22", "mx24", "hpr", "f4p", "atlasti", "exe", "dll", "bin",
}

# Files on disk for these statuses (downloaded or unzipped).
ON_DISK_STATUSES = ("success", "skipped_exists", "extracted")

# Returned by _extract when a file simply isn't text (no parser / sniffed binary),
# as opposed to None which means a parser was tried and errored.
_NOT_TEXT = object()


class TextExtractor:
    def __init__(
        self,
        db: QDArchDatabase,
        source_root: str = "downloads",
        output_root: str = "extracted_text",
        max_file_size_mb: float = 200.0,
        overwrite: bool = False,
    ):
        self.db = db
        self.source_root = Path(source_root)
        self.output_root = Path(output_root)
        self.max_bytes = int(max_file_size_mb * 1024 * 1024)
        self.overwrite = overwrite
        self.output_root.mkdir(parents=True, exist_ok=True)

    # ── Public API ──────────────────────────────────────────────────────────

    def extract_all(self, statuses=ON_DISK_STATUSES) -> dict:
        placeholders = ",".join("?" for _ in statuses)
        files = self.db.query(
            f"SELECT * FROM files WHERE status IN ({placeholders})", tuple(statuses)
        )
        summary = {
            "extracted": 0, "empty": 0, "skipped_type": 0,
            "missing": 0, "failed": 0, "already": 0,
        }
        if files.empty:
            print("No on-disk files to process.")
            return summary

        done = set() if self.overwrite else self.db.get_extracted_file_ids()
        projects: dict[int, object] = {}

        print(f"Scanning {len(files)} on-disk file records...")
        for _, frow in files.iterrows():
            file_id = int(frow["id"])
            if file_id in done:
                summary["already"] += 1
                continue

            ext = _ext_of(frow["file_name"], frow["file_type"])
            if ext in SKIP_EXTS:
                summary["skipped_type"] += 1
                continue

            project_id = int(frow["project_id"])
            if project_id not in projects:
                pdf = self.db.get_project(project_id)
                projects[project_id] = None if pdf.empty else pdf.iloc[0]
            project = projects[project_id]
            if project is None:
                summary["missing"] += 1
                continue

            src_path = self._resolve_source(project, frow["file_name"])
            if not src_path.exists():
                summary["missing"] += 1
                continue

            if src_path.stat().st_size > self.max_bytes:
                print(f"  [skip] {src_path.name} — exceeds {self.max_bytes // 1024 // 1024} MB")
                summary["skipped_type"] += 1
                continue

            text = self._extract(src_path, ext)
            if text is _NOT_TEXT:
                summary["skipped_type"] += 1
                continue
            if text is None:
                summary["failed"] += 1
                continue

            text = text.strip()
            if not text:
                summary["empty"] += 1
                continue

            out_path = self._write_text(project, file_id, frow["file_name"], text)
            self.db.insert_extracted_text(
                file_id=file_id,
                project_id=project_id,
                extracted_path=str(out_path),
                source_path=str(src_path),
                file_type=ext,
                char_count=len(text),
                status="extracted",
            )
            summary["extracted"] += 1
            print(f"  [ok] {src_path.name} -> {out_path.name} ({len(text):,} chars)")

        return summary

    # ── Path handling ─────────────────────────────────────────────────────────

    def _project_subpath(self, project) -> Path:
        p = Path(project["download_repository_folder"]) / str(project["download_project_folder"])
        if project["download_version_folder"]:
            p = p / project["download_version_folder"]
        return p

    def _resolve_source(self, project, file_name: str) -> Path:
        """Locate the file on disk, sanitizing each path component as the
        downloader did when it wrote the file."""
        rel = Path(file_name)
        safe_parts = [_sanitize_path_part(part) for part in rel.parts]
        return self.source_root / self._project_subpath(project) / Path(*safe_parts)

    def _write_text(self, project, file_id: int, file_name: str, text: str) -> Path:
        out_dir = self.output_root / self._project_subpath(project)
        out_dir.mkdir(parents=True, exist_ok=True)
        stem = _sanitize_path_part(Path(file_name).stem) or "file"
        out_path = out_dir / f"{file_id}__{stem}.txt"
        out_path.write_text(text, encoding="utf-8")
        return out_path

    # ── Parsing dispatch ──────────────────────────────────────────────────────

    def _extract(self, path: Path, ext: str) -> Optional[str]:
        try:
            if ext in PDF_EXTS:
                return _parse_pdf(path)
            if ext in DOCX_EXTS:
                return _parse_docx(path)
            if ext in XLSX_EXTS:
                return _parse_xlsx(path)
            if ext in MARKUP_EXTS:
                return _parse_markup(path)
            if ext in PLAIN_EXTS:
                return _parse_plain(path)
            # Unknown extension: sniff the bytes and decode if it looks textual.
            data = path.read_bytes()
            if _looks_textual(data):
                return _decode(data)
            return _NOT_TEXT  # binary, no parser — not a failure
        except Exception as exc:
            log.warning("Failed to extract %s (%s): %s", path, ext, exc)
            print(f"  [fail] {path.name} — {exc}")
            return None


# ── Module-level parsers / helpers ──────────────────────────────────────────

def _ext_of(file_name: str, file_type: Optional[str]) -> str:
    ft = (file_type or "").lower().lstrip(".")
    if ft:
        return ft
    name = Path(file_name or "").name
    return name.rsplit(".", 1)[-1].lower() if "." in name else ""


def _decode(data: bytes) -> str:
    for enc in ("utf-8-sig", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace")


def _looks_textual(data: bytes) -> bool:
    if not data:
        return False
    sample = data[:8192]
    if b"\x00" in sample:
        return False
    printable = sum(1 for b in sample if b in (9, 10, 13) or 32 <= b <= 126 or b >= 128)
    return printable / len(sample) > 0.85


def _parse_pdf(path: Path) -> Optional[str]:
    import fitz  # PyMuPDF
    parts = []
    with fitz.open(path) as doc:
        for page in doc:
            parts.append(page.get_text())
    return "\n".join(parts)


def _parse_docx(path: Path) -> Optional[str]:
    import docx
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _parse_xlsx(path: Path) -> Optional[str]:
    import openpyxl
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    try:
        parts = []
        for ws in wb.worksheets:
            parts.append(f"# Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                parts.append("\t".join("" if c is None else str(c) for c in row))
        return "\n".join(parts)
    finally:
        wb.close()


def _parse_markup(path: Path) -> Optional[str]:
    from lxml import etree, html as lhtml
    data = path.read_bytes()
    try:
        tree = lhtml.fromstring(data)
        return tree.text_content()
    except Exception:
        try:
            root = etree.fromstring(data)
            return " ".join(root.itertext())
        except Exception:
            return _decode(data)


def _parse_plain(path: Path) -> Optional[str]:
    return _decode(path.read_bytes())
